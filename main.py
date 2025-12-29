import os
import io
import uuid
import base64
import json
import logging
import re
from datetime import datetime
from typing import Optional
from urllib.parse import quote_plus

import requests
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import socketio
from dotenv import load_dotenv
from groq import Groq

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# Database imports
from database import SessionLocal, init_db
import crud

# MCP Integration
from mcp_integration import mcp_router, mcp_context

# PDF imports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
load_dotenv()

# Initialize database (SQLite)
init_db()

# Env
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
if not GROQ_API_KEY:
    raise ValueError("❌ GROQ_API_KEY not found!")

GROQ_MODEL = os.getenv("GROQ_MODEL", "").strip() or "llama-3.3-70b-versatile"
groq_client = Groq(api_key=GROQ_API_KEY)

# FastAPI
app = FastAPI(title="NCERT Lesson Plan Generator with MCP")
app.include_router(mcp_router)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Socket.IO (ASGI)
sio = socketio.AsyncServer(
    async_mode="asgi",
    cors_allowed_origins="*",
    logger=True,
    engineio_logger=True,
)
socket_app = socketio.ASGIApp(sio, app)

# In-memory store
lessons_store = {}
conversation_states = {}

# ============================================================
# Links
# ============================================================

def generate_cbse_youtube_links(topic: str, subject: str, class_level: str) -> list:
    """Generate ONLY topic-specific CBSE/NCERT-related YouTube search links."""
    class_num = class_level.replace("Class ", "").strip()

    topic_clean = topic.strip()
    subject_clean = subject.strip()
    youtube_links = []

    cbse_search = quote_plus(f"{topic_clean} class {class_num} {subject_clean}")
    youtube_links.append({
        "title": f"{topic_clean} - CBSE Official",
        "url": f"https://www.youtube.com/@cbabordsecondaryedu/search?query={cbse_search}",
        "description": f"Official CBSE videos for '{topic_clean}' Class {class_num}",
        "channel": "CBSE Official",
    })

    ncert_search = quote_plus(f"{topic_clean} {subject_clean} class {class_num}")
    youtube_links.append({
        "title": f"{topic_clean} - NCERT Official",
        "url": f"https://www.youtube.com/@NCERTOfficial/search?query={ncert_search}",
        "description": f"NCERT official videos for '{topic_clean}'",
        "channel": "NCERT Official",
    })

    diksha_search = quote_plus(f"diksha {topic_clean} class {class_num} {subject_clean} CBSE")
    youtube_links.append({
        "title": f"{topic_clean} - DIKSHA",
        "url": f"https://www.youtube.com/results?search_query={diksha_search}",
        "description": f"DIKSHA educational content for '{topic_clean}' Class {class_num}",
        "channel": "DIKSHA",
    })

    swayam_search = quote_plus(f"swayam prabha {topic_clean} {subject_clean} class {class_num}")
    youtube_links.append({
        "title": f"{topic_clean} - Swayam Prabha",
        "url": f"https://www.youtube.com/results?search_query={swayam_search}",
        "description": f"Swayam Prabha educational videos for '{topic_clean}'",
        "channel": "Swayam Prabha",
    })

    return youtube_links


def generate_ncert_web_resources(topic: str, subject: str, class_level: str) -> list:
    """Generate NCERT and educational portal links."""
    class_num = class_level.replace("Class ", "").strip()
    topic_encoded = quote_plus(topic)

    return [
        {
            "title": f"NCERT Textbook - Class {class_num} {subject}",
            "url": "https://ncert.nic.in/textbook.php",
            "description": f"Official NCERT textbook for Class {class_num} {subject}",
            "type": "NCERT Textbook",
        },
        {
            "title": "e-Pathshala - Digital Textbooks",
            "url": "https://epathshala.nic.in/",
            "description": f"Digital NCERT textbooks and resources for {subject}",
            "type": "e-Pathshala",
        },
        {
            "title": f"DIKSHA - {topic}",
            "url": f"https://diksha.gov.in/explore?searchQuery={topic_encoded}&board=CBSE&gradeLevel=Class%20{class_num}",
            "description": f"Interactive learning content for '{topic}' on DIKSHA",
            "type": "DIKSHA Portal",
        },
        {
            "title": f"NROER - {topic}",
            "url": f"https://nroer.gov.in/home/search/?search_text={topic_encoded}",
            "description": f"Open educational resources for '{topic}'",
            "type": "NROER",
        },
    ]


# ============================================================
# Conversation State
# ============================================================

class ConversationState:
    def __init__(self, sid: str):
        self.sid = sid
        self.stage = "LANGUAGE_SELECTION"
        self.language_mode: Optional[str] = None  # "english" / "hindi"
        self.data = {
            "topic": None,
            "subject": None,
            "class_level": None,
            "session_duration": 40,
            "num_sessions": 4,
            "language": "English",
        }
        self.history = []

    def add_to_history(self, role: str, text: str):
        self.history.append({"role": role, "text": text})


PROMPTS = {
    "LANGUAGE_SELECTION": {
        "english": (
            "Hello! I'm your NCERT Lesson Plan Assistant.\n\n"
            "I can help you create a complete lesson plan through voice conversation.\n\n"
            "Which language would you prefer?\n"
            "1) English\n"
            "2) हिंदी (Hindi)\n\n"
            "Please say 'English' or 'Hindi'."
        ),
        "hindi": (
            "नमस्ते! मैं आपका NCERT पाठ योजना सहायक हूँ।\n\n"
            "मैं आवाज़ के माध्यम से एक पूर्ण पाठ योजना बनाने में आपकी सहायता कर सकता हूँ।\n\n"
            "आप किस भाषा को पसंद करेंगे?\n"
            "1) English\n"
            "2) हिंदी (Hindi)\n\n"
            "कृपया 'English' या 'Hindi' कहें।"
        ),
    },
    "TOPIC_COLLECTION": {
        "english": (
            "Great! Let's start creating your lesson plan.\n\n"
            "What is the topic or lesson title?\n"
            "Example: Photosynthesis, Fractions, Freedom Struggle"
        ),
        "hindi": (
            "बढ़िया! चलिए आपकी पाठ योजना बनाना शुरू करते हैं।\n\n"
            "पाठ का विषय या शीर्षक क्या है?\n"
            "उदाहरण: प्रकाश संश्लेषण, भिन्न, स्वतंत्रता संग्राम"
        ),
    },
    "SUBJECT_COLLECTION": {
        "english": "Which subject is this lesson for? Example: Mathematics, Science, Social Science, English, Hindi.",
        "hindi": "यह पाठ किस विषय के लिए है? उदाहरण: गणित, विज्ञान, सामाजिक विज्ञान, अंग्रेज़ी, हिंदी।",
    },
    "CLASS_COLLECTION": {
        "english": "Which class is this lesson for? Please say class number 1 to 12. Example: Class 5.",
        "hindi": "यह पाठ किस कक्षा के लिए है? कृपया 1 से 12 के बीच कक्षा संख्या बताएं।",
    },
    "SESSION_DURATION": {
        "english": "How long should each session be? Please say 15 to 90 minutes. Example: 40 minutes.",
        "hindi": "प्रत्येक सत्र कितने मिनट का होना चाहिए? 15 से 90 मिनट। उदाहरण: 40 मिनट।",
    },
    "NUM_SESSIONS": {
        "english": "How many sessions do you need? Please say 1 to 10. Example: 4 sessions.",
        "hindi": "आपको कितने सत्र चाहिए? 1 से 10। उदाहरण: 4 सत्र।",
    },
    "CONFIRMATION": {
        "english": lambda d: (
            "Perfect! Please confirm:\n\n"
            f"Topic: {d['topic']}\n"
            f"Subject: {d['subject']}\n"
            f"Class: {d['class_level']}\n"
            f"Session Duration: {d['session_duration']} minutes\n"
            f"Number of Sessions: {d['num_sessions']}\n\n"
            "Say 'Yes' to generate, or 'No' to start over."
        ),
        "hindi": lambda d: (
            "कृपया पुष्टि करें:\n\n"
            f"विषय: {d['topic']}\n"
            f"सब्जेक्ट: {d['subject']}\n"
            f"कक्षा: {d['class_level']}\n"
            f"सत्र अवधि: {d['session_duration']} मिनट\n"
            f"सत्रों की संख्या: {d['num_sessions']}\n\n"
            "बनाने के लिए 'हाँ' कहें, या फिर से शुरू करने के लिए 'नहीं' कहें।"
        ),
    },
}


def get_db():
    db = SessionLocal()
    try:
        return db
    finally:
        pass


# ============================================================
# Groq STT + Chat
# ============================================================

def transcribe_audio_with_groq(audio_bytes: bytes, filename: str = "recording.webm") -> str:
    url = "https://api.groq.com/openai/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files = {"file": (filename, audio_bytes, "audio/webm")}
    data = {"model": "whisper-large-v3", "response_format": "json"}
    r = requests.post(url, headers=headers, files=files, data=data, timeout=120)
    if not r.ok:
        raise RuntimeError(f"Groq STT failed: {r.status_code} {r.text}")
    return (r.json().get("text") or "").strip()


def groq_chat(prompt: str, system: str, max_tokens: int = 900, temperature: float = 0.4) -> str:
    resp = groq_client.chat.completions.create(
        model=GROQ_MODEL,
        temperature=temperature,
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ],
    )
    return (resp.choices[0].message.content or "").strip()


# ============================================================
# Parsing user speech
# ============================================================

def parse_user_response(transcript: str, stage: str, language_mode: Optional[str]) -> dict:
    transcript_lower = transcript.lower().strip()

    if any(cmd in transcript_lower for cmd in ["start over", "restart", "शुरू करो", "फिर से"]):
        return {"value": "RESTART", "needs_clarification": False}

    if stage == "LANGUAGE_SELECTION":
        if "english" in transcript_lower or "इंग्लिश" in transcript_lower:
            return {"value": "english", "needs_clarification": False}
        if "hindi" in transcript_lower or "हिंदी" in transcript_lower:
            return {"value": "hindi", "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    if stage == "TOPIC_COLLECTION":
        return {"value": transcript.strip(), "needs_clarification": len(transcript.strip()) < 3}

    if stage == "SUBJECT_COLLECTION":
        subjects_map = {
            "math": "Mathematics", "maths": "Mathematics", "गणित": "Mathematics",
            "science": "Science", "विज्ञान": "Science",
            "english": "English", "अंग्रेजी": "English",
            "hindi": "Hindi", "हिंदी": "Hindi",
            "social": "Social Science", "सामाजिक": "Social Science",
            "history": "History", "इतिहास": "History",
            "geography": "Geography", "भूगोल": "Geography",
            "physics": "Physics", "भौतिकी": "Physics",
            "chemistry": "Chemistry", "रसायन": "Chemistry",
            "biology": "Biology", "जीव विज्ञान": "Biology",
            "computer": "Computer Science", "कंप्यूटर": "Computer Science",
        }
        for key, value in subjects_map.items():
            if key in transcript_lower:
                return {"value": value, "needs_clarification": False}
        if len(transcript.strip()) > 2:
            return {"value": transcript.strip().title(), "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    if stage == "CLASS_COLLECTION":
        numbers = re.findall(r"\d+", transcript)
        if numbers:
            class_num = int(numbers[0])
            if 1 <= class_num <= 12:
                return {"value": f"Class {class_num}", "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    if stage == "SESSION_DURATION":
        numbers = re.findall(r"\d+", transcript)
        if numbers:
            duration = int(numbers[0])
            if 15 <= duration <= 90:
                return {"value": duration, "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    if stage == "NUM_SESSIONS":
        numbers = re.findall(r"\d+", transcript)
        if numbers:
            num = int(numbers[0])
            if 1 <= num <= 10:
                return {"value": num, "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    if stage == "CONFIRMATION":
        yes_words = ["yes", "yeah", "correct", "right", "okay", "ok", "sure", "हां", "हाँ", "सही", "ठीक", "जी"]
        no_words = ["no", "नहीं", "गलत", "wrong", "नही"]
        if any(w in transcript_lower for w in yes_words):
            return {"value": True, "needs_clarification": False}
        if any(w in transcript_lower for w in no_words):
            return {"value": False, "needs_clarification": False}
        return {"value": None, "needs_clarification": True}

    return {"value": None, "needs_clarification": True}


# ============================================================
# ENHANCED Worksheet Generation
# ============================================================

def generate_worksheets(topic: str, subject: str, class_level: str, language: str, num_sessions: int = 4) -> list:
    """Generate detailed worksheets for each session with structured format."""
    class_num = class_level.replace("Class ", "").strip()

    prompt = f"""Create {num_sessions} comprehensive and detailed worksheets for teachers to use in class.

Topic: {topic}
Subject: {subject}
Class: {class_num}
Language: {language}

For EACH worksheet, provide the following EXACT structure:

===WORKSHEET 1===
TITLE: [Specific title related to {topic}]
SESSION: 1
OBJECTIVE: [Clear learning objective for this worksheet]
DURATION: [Suggested completion time: 15-20 minutes]

SECTION A: FILL IN THE BLANKS (5 questions)
Instructions: Complete the sentences with appropriate words.
1. ______________________
2. ______________________
3. ______________________
4. ______________________
5. ______________________

SECTION B: TRUE OR FALSE (5 questions)
Instructions: Write True or False for each statement.
1. ______________________ (True/False)
2. ______________________ (True/False)
3. ______________________ (True/False)
4. ______________________ (True/False)
5. ______________________ (True/False)

SECTION C: MULTIPLE CHOICE QUESTIONS (4 questions)
Instructions: Choose the correct answer.
1. ______________________
   a) ______ b) ______ c) ______ d) ______
2. ______________________
   a) ______ b) ______ c) ______ d) ______
3. ______________________
   a) ______ b) ______ c) ______ d) ______
4. ______________________
   a) ______ b) ______ c) ______ d) ______

SECTION D: SHORT ANSWER QUESTIONS (3 questions)
Instructions: Answer in 2-3 sentences.
1. ______________________
2. ______________________
3. ______________________

SECTION E: ACTIVITY/PRACTICAL TASK (1-2 activities)
Instructions: Complete the following hands-on activities.
1. ______________________
2. ______________________

ANSWER KEY:
Section A Answers:
1. ______ 2. ______ 3. ______ 4. ______ 5. ______

Section B Answers:
1. ______ 2. ______ 3. ______ 4. ______ 5. ______

Section C Answers:
1. ______ 2. ______ 3. ______ 4. ______

Section D Sample Answers:
1. ______________________
2. ______________________
3. ______________________

===END WORKSHEET 1===

Generate all {num_sessions} worksheets following this exact format. Each worksheet should:
- Be progressively more challenging
- Cover different aspects of '{topic}'
- Include age-appropriate questions for Class {class_num}
- Have clear instructions for each section
- Include complete answer keys
"""

    worksheet_text = groq_chat(
        prompt=prompt,
        system=f"You are an expert {subject} curriculum designer creating comprehensive worksheets for Class {class_num} students. Generate detailed, educational, and engaging worksheets with proper formatting.",
        max_tokens=4000,
        temperature=0.5,
    )

    worksheets = parse_worksheets(worksheet_text, topic, num_sessions)
    return worksheets


def parse_worksheets(worksheet_text: str, topic: str, num_sessions: int) -> list:
    """Parse the generated worksheet text into structured format."""
    worksheets = []
    
    # Try to split by worksheet markers
    patterns = [
        r"===WORKSHEET\s+(\d+)===",
        r"---WORKSHEET\s+(\d+)---",
        r"WORKSHEET\s+(\d+):",
        r"\*\*WORKSHEET\s+(\d+)\*\*",
    ]
    
    sections = None
    for pattern in patterns:
        sections = re.split(pattern, worksheet_text, flags=re.IGNORECASE)
        if len(sections) > 1:
            break
    
    if sections and len(sections) > 1:
        # Process paired sections (number, content)
        i = 1
        while i < len(sections):
            if i + 1 < len(sections):
                ws_num = sections[i].strip()
                ws_content = sections[i + 1].strip()
                
                # Clean up content - remove end markers
                ws_content = re.sub(r"===END WORKSHEET \d+===", "", ws_content, flags=re.IGNORECASE)
                ws_content = re.sub(r"---END WORKSHEET \d+---", "", ws_content, flags=re.IGNORECASE)
                
                # Extract title
                title_match = re.search(r"TITLE:\s*(.+?)(?:\n|SESSION:)", ws_content, re.IGNORECASE)
                title = title_match.group(1).strip() if title_match else f"Worksheet {ws_num}: {topic}"
                
                # Extract objective
                obj_match = re.search(r"OBJECTIVE:\s*(.+?)(?:\n|DURATION:)", ws_content, re.IGNORECASE)
                objective = obj_match.group(1).strip() if obj_match else ""
                
                # Extract duration
                dur_match = re.search(r"DURATION:\s*(.+?)(?:\n|SECTION)", ws_content, re.IGNORECASE)
                duration = dur_match.group(1).strip() if dur_match else "20 minutes"
                
                # Extract sections
                sections_data = extract_worksheet_sections(ws_content)
                
                worksheets.append({
                    "number": int(ws_num) if ws_num.isdigit() else len(worksheets) + 1,
                    "title": title,
                    "objective": objective,
                    "duration": duration,
                    "content": ws_content,
                    "sections": sections_data,
                })
                i += 2
            else:
                i += 1
    
    # Fallback: if parsing failed, create basic worksheets
    if not worksheets:
        # Try to find any structured content
        content_chunks = worksheet_text.split("\n\n")
        chunk_size = len(content_chunks) // num_sessions if len(content_chunks) >= num_sessions else 1
        
        for i in range(num_sessions):
            start_idx = i * chunk_size
            end_idx = start_idx + chunk_size if i < num_sessions - 1 else len(content_chunks)
            chunk_content = "\n\n".join(content_chunks[start_idx:end_idx])
            
            worksheets.append({
                "number": i + 1,
                "title": f"Worksheet {i + 1}: {topic}",
                "objective": f"Practice and reinforce understanding of {topic}",
                "duration": "20 minutes",
                "content": chunk_content if chunk_content.strip() else worksheet_text,
                "sections": {},
            })
    
    return worksheets


def extract_worksheet_sections(content: str) -> dict:
    """Extract individual sections from worksheet content."""
    sections = {}
    
    section_patterns = {
        "fill_blanks": r"SECTION A[:\s]*FILL IN THE BLANKS.*?(?=SECTION B|$)",
        "true_false": r"SECTION B[:\s]*TRUE OR FALSE.*?(?=SECTION C|$)",
        "mcq": r"SECTION C[:\s]*MULTIPLE CHOICE.*?(?=SECTION D|$)",
        "short_answer": r"SECTION D[:\s]*SHORT ANSWER.*?(?=SECTION E|ANSWER KEY|$)",
        "activity": r"SECTION E[:\s]*ACTIVITY.*?(?=ANSWER KEY|$)",
        "answer_key": r"ANSWER KEY:.*?(?=={3}|$)",
    }
    
    for section_name, pattern in section_patterns.items():
        match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if match:
            sections[section_name] = match.group(0).strip()
    
    return sections


# ============================================================
# Generation helpers
# ============================================================

def generate_lesson_from_topic(topic: str, subject: str, class_level: str, language: str,
                              session_duration: int = 40, num_sessions: int = 4) -> str:
    today = datetime.now().strftime("%d/%m/%Y")
    total_duration = session_duration * num_sessions

    prompt = f"""Create a complete NCERT lesson plan for the given topic.

Topic: {topic}
Subject: {subject}
Class: {class_level}
Language: {language}
Number of Sessions: {num_sessions}
Duration per Session: {session_duration} minutes
Total Duration: {total_duration} minutes

Generate in EXACT structure:

LESSON PLAN
Subject: {subject}
Class: {class_level}
Topic: {topic}
Duration: {total_duration} minutes ({num_sessions} sessions × {session_duration} mins each)
Date: {today}

---LEARNING OBJECTIVES---
(3-4 bullet points)

---LEARNING OUTCOMES---
(3-4 bullet points)

---PRE-REQUISITE KNOWLEDGE---
(2-3 bullet points)

---TEACHING AIDS/RESOURCES---
(list)

---INTRODUCTION---
(5-8 lines)

---MAIN CONTENT---
(8-15 lines)

---ACTIVITIES---
({num_sessions * 2}-{num_sessions * 3} bullet points)

---ASSESSMENT---
(4-6 bullet points)

---HOMEWORK---
(2-4 bullet points)

---CONCLUSION---
(4-6 lines)

---REFLECTION---
(4-6 bullet points)
"""
    return groq_chat(
        prompt=prompt,
        system="You are an expert NCERT teacher who writes complete lesson plans.",
        max_tokens=2500,
        temperature=0.5,
    )


def _extract_section(full: str, section_name: str) -> str:
    marker = f"---{section_name}---"
    if marker not in full:
        return ""
    after = full.split(marker, 1)[1]
    next_idx = after.find("\n---")
    return after.strip() if next_idx == -1 else after[:next_idx].strip()


def _lines_to_list(block: str) -> list:
    if not block:
        return []
    out = []
    for raw in block.splitlines():
        s = raw.strip()
        if not s:
            continue
        s = s.lstrip("•- \t")
        if len(s) > 2 and s[0].isdigit() and s[1] in [".", ")"]:
            s = s[2:].strip()
        out.append(s)
    return out


def _split_into_n(items: list, n: int) -> list:
    if not items:
        return [["-"] for _ in range(n)]
    buckets = [[] for _ in range(n)]
    for idx, item in enumerate(items):
        buckets[idx % n].append(item)
    return [b if b else ["-"] for b in buckets]


def parse_llm_to_sessions(formatted_text: str, session_duration: int = 40, num_sessions: int = 4) -> list:
    objectives = _extract_section(formatted_text, "LEARNING OBJECTIVES")
    outcomes = _extract_section(formatted_text, "LEARNING OUTCOMES")
    resources = _extract_section(formatted_text, "TEACHING AIDS/RESOURCES")
    activities = _extract_section(formatted_text, "ACTIVITIES")
    assessment = _extract_section(formatted_text, "ASSESSMENT")
    homework = _extract_section(formatted_text, "HOMEWORK")

    obj_list = _lines_to_list(objectives)
    out_list = _lines_to_list(outcomes)
    res_list = _lines_to_list(resources)
    act_list = _lines_to_list(activities)
    assess_list = _lines_to_list(assessment)
    hw_list = _lines_to_list(homework)

    act_n = _split_into_n(act_list, num_sessions)
    res_n = _split_into_n(res_list, num_sessions)
    assess_n = _split_into_n(assess_list, num_sessions)

    def pick(lst, i, fallback):
        return lst[i % len(lst)] if lst else fallback

    sessions = []
    for i in range(num_sessions):
        sessions.append({
            "number": i + 1,
            "duration": f"{session_duration} mins",
            "competency": pick(obj_list, i, "Competency based on objectives"),
            "elo": pick(out_list, i, "Expected learning outcome"),
            "activities": act_n[i],
            "resources_tlm": "; ".join(res_n[i]) if res_n[i] != ["-"] else "-",
            "e_resources": [],
            "worksheets": f"Worksheet {i+1}",
            "assessment": "; ".join(assess_n[i]) if assess_n[i] != ["-"] else ("; ".join(hw_list) if hw_list else "-"),
        })
    return sessions


def build_lesson_data_from_formatted(topic: str, subject: str, class_level: str, language: str,
                                    formatted: str, session_duration: int = 40, num_sessions: int = 4) -> dict:
    class_num = class_level.replace("Class ", "").strip()
    
    # Generate comprehensive worksheets
    worksheets = generate_worksheets(topic, subject, class_level, language, num_sessions)
    youtube_links = generate_cbse_youtube_links(topic, subject, class_level)
    web_resources = generate_ncert_web_resources(topic, subject, class_level)

    # Parse sessions
    sessions = parse_llm_to_sessions(formatted, session_duration, num_sessions)
    
    # Link worksheets to sessions
    for i, session in enumerate(sessions):
        if i < len(worksheets):
            session["worksheets"] = worksheets[i].get("title", f"Worksheet {i+1}")
            session["worksheet_data"] = worksheets[i]

    return {
        "header": {
            "class": class_num,
            "subject": subject,
            "lesson": topic,
            "periods": str(num_sessions),
            "duration": f"{session_duration} mins",
            "total_duration": f"{session_duration * num_sessions} mins",
            "language": language,
        },
        "sessions": sessions,
        "formatted_text": formatted,
        "worksheets": worksheets,
        "youtube_links": youtube_links,
        "web_resources": web_resources,
    }


# ============================================================
# ENHANCED DOCX with Worksheets
# ============================================================

def create_docx_lesson_plan(lesson_data: dict, topic: str) -> io.BytesIO:
    doc = Document()

    # Title
    title = doc.add_heading("NCERT LESSON PLANNER", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading("LESSON PLAN", 0)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    num_sessions = len(lesson_data["sessions"])
    session_dur = lesson_data["header"].get("duration", "40 mins")
    total_dur = lesson_data["header"].get("total_duration", f"{num_sessions * 40} mins")

    # Header table
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"

    details = [
        ("Class", lesson_data["header"]["class"]),
        ("Subject", lesson_data["header"]["subject"]),
        ("Lesson", lesson_data["header"]["lesson"]),
        ("No. of Periods/Sessions", lesson_data["header"]["periods"]),
        ("Session Duration", session_dur),
        ("Total Duration", total_dur),
        ("Language", lesson_data["header"].get("language", "English")),
    ]
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = f"{label}:"
        table.rows[i].cells[1].text = str(value)
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph(f"\nThis lesson needs {num_sessions} Sessions to complete.\n")

    # Sessions table
    headers = [
        "Session (Duration)", "Competency", "ELO", "Suggested Activities",
        "Suggested Resources/TLM", "E-Resources", "Worksheets", "Assessment"
    ]
    session_table = doc.add_table(rows=1, cols=len(headers))
    session_table.style = "Table Grid"
    hdr_cells = session_table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        hdr_cells[j].paragraphs[0].runs[0].bold = True

    for s in lesson_data["sessions"]:
        row = session_table.add_row().cells
        row[0].text = f"Session {s['number']} ({s.get('duration', '')})"
        row[1].text = s.get("competency", "-")
        row[2].text = s.get("elo", "-")
        acts = s.get("activities", [])
        row[3].text = "\n".join([f"• {a}" for a in acts if a and a != "-"]) or "-"
        row[4].text = s.get("resources_tlm", "-")
        row[5].text = "\n".join(s.get("e_resources", [])) or "-"
        row[6].text = s.get("worksheets", "-")
        row[7].text = s.get("assessment", "-")

    # ============================================================
    # WORKSHEETS SECTION - ENHANCED
    # ============================================================
    doc.add_page_break()
    ws_title = doc.add_heading("TEACHER WORKSHEETS", 1)
    ws_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "The following worksheets are designed to complement each session of the lesson plan. "
        "Each worksheet includes various question types and an answer key for teacher reference."
    )
    doc.add_paragraph("")

    for worksheet in lesson_data.get("worksheets", []):
        # Worksheet header
        ws_heading = doc.add_heading(f"", 2)
        ws_run = ws_heading.add_run(worksheet.get("title", f"Worksheet {worksheet.get('number', '')}"))
        ws_run.bold = True
        
        # Worksheet metadata table
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.style = "Light List"
        
        meta_data = [
            ("Objective:", worksheet.get("objective", "Practice and assessment")),
            ("Duration:", worksheet.get("duration", "20 minutes")),
            ("Session:", f"Session {worksheet.get('number', '')}"),
        ]
        
        for i, (label, value) in enumerate(meta_data):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            meta_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph("")
        
        # Worksheet content
        content = worksheet.get("content", "") or ""
        
        # Process content line by line with proper formatting
        current_section = None
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            # Section headers
            if re.match(r"^SECTION [A-E]", line, re.IGNORECASE):
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 51, 102)
                current_section = line
                
            # Instructions
            elif line.lower().startswith("instructions:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.size = Pt(10)
                
            # Title, Objective, Duration labels
            elif any(line.startswith(label) for label in ["TITLE:", "OBJECTIVE:", "DURATION:", "SESSION:"]):
                continue  # Already handled in metadata table
                
            # Answer Key header
            elif "ANSWER KEY" in line.upper():
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run("📝 ANSWER KEY")
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 102, 51)
                
            # Numbered questions
            elif re.match(r"^\d+[\.\)]\s*", line):
                p = doc.add_paragraph(style='List Number')
                # Clean the number prefix
                clean_line = re.sub(r"^\d+[\.\)]\s*", "", line)
                p.add_run(clean_line)
                
            # Multiple choice options
            elif re.match(r"^[a-d]\)", line, re.IGNORECASE):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(line)
                
            # Section answer labels
            elif re.match(r"^Section [A-E] Answers?:", line, re.IGNORECASE):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
                
            # Regular content
            else:
                doc.add_paragraph(line)
        
        # Separator between worksheets
        doc.add_paragraph("")
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.add_run("─" * 60)
        doc.add_paragraph("")

    # ============================================================
    # YouTube Links Section
    # ============================================================
    doc.add_page_break()
    doc.add_heading("CBSE OFFICIAL YOUTUBE RESOURCES", 1)
    doc.add_paragraph(f"Topic-specific video resources for '{topic}' from official channels:")

    yt_table = doc.add_table(rows=1, cols=3)
    yt_table.style = "Table Grid"
    yt_hdr = yt_table.rows[0].cells
    yt_hdr[0].text = "Title"
    yt_hdr[1].text = "Description"
    yt_hdr[2].text = "URL"
    for cell in yt_hdr:
        cell.paragraphs[0].runs[0].bold = True

    for link in lesson_data.get("youtube_links", []):
        row = yt_table.add_row().cells
        row[0].text = link.get("title", "-")
        row[1].text = link.get("description", "-")
        row[2].text = link.get("url", "-")

    # Web resources
    doc.add_paragraph("")
    doc.add_heading("NCERT WEB RESOURCES", 2)
    web_table = doc.add_table(rows=1, cols=3)
    web_table.style = "Table Grid"
    web_hdr = web_table.rows[0].cells
    web_hdr[0].text = "Resource"
    web_hdr[1].text = "Description"
    web_hdr[2].text = "URL"
    for cell in web_hdr:
        cell.paragraphs[0].runs[0].bold = True

    for link in lesson_data.get("web_resources", []):
        row = web_table.add_row().cells
        row[0].text = link.get("title", "-")
        row[1].text = link.get("description", "-")
        row[2].text = link.get("url", "-")

    # Footer
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"Generated on: {datetime.now().strftime('%d %B %Y at %H:%M')}")
    run.italic = True
    run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# ENHANCED PDF with Worksheets
# ============================================================

def create_pdf_lesson_plan(lesson_data: dict, topic: str) -> io.BytesIO:
    if not PDF_AVAILABLE:
        raise Exception("ReportLab not installed")

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        title=f"NCERT Lesson Plan - {topic}",
    )

    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        "NCERT_title",
        parent=styles["Title"],
        alignment=1,
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1e40af"),
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "NCERT_heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#1e40af"),
        spaceBefore=15,
        spaceAfter=8,
    )
    subheading_style = ParagraphStyle(
        "NCERT_subheading",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#374151"),
        spaceBefore=10,
        spaceAfter=6,
    )
    normal_style = ParagraphStyle(
        "NCERT_normal",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        spaceAfter=6,
    )
    small_style = ParagraphStyle(
        "NCERT_small",
        parent=styles["Normal"],
        fontSize=8.5,
        leading=10.5,
    )
    small_bold_style = ParagraphStyle(
        "NCERT_small_bold",
        parent=small_style,
        fontName="Helvetica-Bold",
        alignment=1,
    )
    worksheet_title_style = ParagraphStyle(
        "WS_title",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#065f46"),
        spaceBefore=12,
        spaceAfter=6,
    )
    section_header_style = ParagraphStyle(
        "WS_section",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#1e40af"),
        spaceBefore=8,
        spaceAfter=4,
    )
    answer_key_style = ParagraphStyle(
        "WS_answer",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#065f46"),
        spaceBefore=10,
        spaceAfter=4,
    )

    story = []
    
    # Title page
    story.append(Paragraph("NCERT LESSON PLANNER", title_style))
    story.append(Paragraph("LESSON PLAN", title_style))
    story.append(Spacer(1, 12))

    num_sessions = len(lesson_data["sessions"])
    session_dur = lesson_data["header"].get("duration", "40 mins")
    total_dur = lesson_data["header"].get("total_duration", f"{num_sessions * 40} mins")

    # Header details table
    details_data = [
        [Paragraph("Class:", normal_style), Paragraph(str(lesson_data["header"]["class"]), normal_style)],
        [Paragraph("Subject:", normal_style), Paragraph(str(lesson_data["header"]["subject"]), normal_style)],
        [Paragraph("Lesson:", normal_style), Paragraph(str(lesson_data["header"]["lesson"]), normal_style)],
        [Paragraph("No. of Sessions:", normal_style), Paragraph(str(num_sessions), normal_style)],
        [Paragraph("Duration per Session:", normal_style), Paragraph(str(session_dur), normal_style)],
        [Paragraph("Total Duration:", normal_style), Paragraph(str(total_dur), normal_style)],
        [Paragraph("Language:", normal_style), Paragraph(str(lesson_data["header"].get("language", "English")), normal_style)],
    ]
    details_table = Table(details_data, colWidths=[2.0 * inch, 5.0 * inch])
    details_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e40af")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e0f2fe")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(details_table)
    story.append(Spacer(1, 12))

    # Sessions table
    headers = ["Session", "Competency", "ELO", "Activities", "Resources", "Worksheets", "Assessment"]
    table_data = [[Paragraph(h, small_bold_style) for h in headers]]

    for s in lesson_data["sessions"]:
        acts = s.get("activities", [])
        acts_text = "<br/>".join([f"• {a}" for a in acts if a and a != "-"]) or "-"
        table_data.append([
            Paragraph(f"Session {s.get('number')} ({s.get('duration', '-')})", small_style),
            Paragraph(s.get("competency", "-"), small_style),
            Paragraph(s.get("elo", "-"), small_style),
            Paragraph(acts_text, small_style),
            Paragraph(s.get("resources_tlm", "-"), small_style),
            Paragraph(s.get("worksheets", "-"), small_style),
            Paragraph(s.get("assessment", "-"), small_style),
        ])

    sessions_table = Table(table_data, colWidths=[70, 85, 85, 140, 80, 60, 80], repeatRows=1)
    sessions_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(sessions_table)

    # ============================================================
    # WORKSHEETS SECTION IN PDF
    # ============================================================
    story.append(PageBreak())
    story.append(Paragraph("TEACHER WORKSHEETS", title_style))
    story.append(Paragraph(
        "The following worksheets complement each session and include various question types with answer keys.",
        normal_style
    ))
    story.append(Spacer(1, 12))

    for worksheet in lesson_data.get("worksheets", []):
        # Worksheet title
        ws_title = worksheet.get("title", f"Worksheet {worksheet.get('number', '')}")
        story.append(Paragraph(f"📝 {ws_title}", worksheet_title_style))
        
        # Worksheet metadata
        meta_data = [
            [Paragraph("<b>Objective:</b>", small_style), 
             Paragraph(worksheet.get("objective", "Practice and assessment"), small_style)],
            [Paragraph("<b>Duration:</b>", small_style), 
             Paragraph(worksheet.get("duration", "20 minutes"), small_style)],
            [Paragraph("<b>Session:</b>", small_style), 
             Paragraph(f"Session {worksheet.get('number', '')}", small_style)],
        ]
        meta_table = Table(meta_data, colWidths=[1.2 * inch, 5.8 * inch])
        meta_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0fdf4")),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#065f46")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 8))
        
        # Worksheet content
        content = worksheet.get("content", "") or ""
        
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            # Skip already-handled metadata
            if any(line.upper().startswith(label) for label in ["TITLE:", "OBJECTIVE:", "DURATION:", "SESSION:"]):
                continue
            
            # Section headers
            if re.match(r"^SECTION [A-E]", line, re.IGNORECASE):
                story.append(Spacer(1, 6))
                story.append(Paragraph(line, section_header_style))
                
            # Instructions
            elif line.lower().startswith("instructions:"):
                story.append(Paragraph(f"<i>{line}</i>", small_style))
                
            # Answer Key header
            elif "ANSWER KEY" in line.upper():
                story.append(Spacer(1, 8))
                story.append(Paragraph("📝 ANSWER KEY", answer_key_style))
                
            # Numbered questions/items
            elif re.match(r"^\d+[\.\)]\s*", line):
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{line}", small_style))
                
            # Multiple choice options
            elif re.match(r"^[a-d]\)", line, re.IGNORECASE):
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{line}", small_style))
                
            # Section answer labels
            elif re.match(r"^Section [A-E] Answers?:", line, re.IGNORECASE):
                story.append(Paragraph(f"<b>{line}</b>", small_style))
                
            # Regular content
            else:
                # Escape special characters for ReportLab
                safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_line, small_style))
        
        # Separator between worksheets
        story.append(Spacer(1, 12))
        story.append(Paragraph("─" * 80, ParagraphStyle("separator", alignment=1, fontSize=8, textColor=colors.grey)))
        story.append(Spacer(1, 12))

    # ============================================================
    # YouTube / Web resources pages
    # ============================================================
    story.append(PageBreak())
    story.append(Paragraph("CBSE OFFICIAL YOUTUBE RESOURCES", heading_style))
    story.append(Paragraph(f"Topic-specific video resources for '{topic}':", normal_style))
    story.append(Spacer(1, 8))

    yt = lesson_data.get("youtube_links", [])
    if yt:
        yt_data = [[Paragraph("Title", small_bold_style), Paragraph("Channel", small_bold_style), Paragraph("URL", small_bold_style)]]
        for link in yt:
            yt_data.append([
                Paragraph(link.get("title", "-"), small_style),
                Paragraph(link.get("channel", "-"), small_style),
                Paragraph(link.get("url", "-"), small_style),
            ])
        yt_table = Table(yt_data, colWidths=[2.3 * inch, 1.3 * inch, 3.4 * inch])
        yt_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(yt_table)

    story.append(Spacer(1, 12))
    story.append(Paragraph("NCERT WEB RESOURCES", heading_style))

    web = lesson_data.get("web_resources", [])
    if web:
        web_data = [[Paragraph("Resource", small_bold_style), Paragraph("Type", small_bold_style), Paragraph("URL", small_bold_style)]]
        for link in web:
            web_data.append([
                Paragraph(link.get("title", "-"), small_style),
                Paragraph(link.get("type", "-"), small_style),
                Paragraph(link.get("url", "-"), small_style),
            ])
        web_table = Table(web_data, colWidths=[2.6 * inch, 1.3 * inch, 3.1 * inch])
        web_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(web_table)

    # Footer
    story.append(Spacer(1, 20))
    footer_style = ParagraphStyle("footer", alignment=1, fontSize=9, textColor=colors.grey, fontName="Helvetica-Oblique")
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y at %H:%M')}", footer_style))

    pdf.build(story)
    buffer.seek(0)
    return buffer


# ============================================================
# FastAPI Routes
# ============================================================

@app.get("/", response_class=HTMLResponse)
async def index():
    with open("templates/index_final.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/mcp/health")
async def mcp_health():
    tools = mcp_context.list_tools()
    return {
        "status": "healthy",
        "mcp_enabled": True,
        "available_tools": len(tools),
        "tools": [t["name"] for t in tools],
    }


@app.post("/download/{lesson_id}")
async def download_lesson(lesson_id: str, request: Request):
    if lesson_id not in lessons_store:
        raise HTTPException(status_code=404, detail="Lesson not found")

    data = lessons_store[lesson_id]
    lesson_data = data["lesson_data"]
    topic = data["topic"]

    try:
        body = await request.json()
    except Exception:
        body = {}

    file_format = (body.get("format") or "docx").lower()

    try:
        if file_format == "pdf":
            buffer = create_pdf_lesson_plan(lesson_data, topic)
            filename = f"NCERT_LessonPlan_{topic.replace(' ', '_')}.pdf"
            media_type = "application/pdf"
        else:
            buffer = create_docx_lesson_plan(lesson_data, topic)
            filename = f"NCERT_LessonPlan_{topic.replace(' ', '_')}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        logging.exception("Download failed")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")


# ============================================================
# Socket.IO Handlers
# ============================================================

@sio.event
async def connect(sid, environ):
    print(f"👤 Client {sid} connected")


@sio.event
async def disconnect(sid):
    if sid in conversation_states:
        del conversation_states[sid]
    print(f"👤 Client {sid} disconnected")


@sio.event
async def textmessage(sid, data):
    topic = (data.get("topic") or "Lesson Topic").strip()
    subject = (data.get("subject") or "Mathematics").strip()
    class_level = (data.get("classlevel") or "Class 2").strip()
    language = (data.get("language") or "English").strip()
    raw_content = (data.get("content") or "").strip()

    session_duration = int(data.get("sessionduration") or 40)
    num_sessions = int(data.get("numsessions") or 4)

    session_duration = max(15, min(90, session_duration))
    num_sessions = max(1, min(10, num_sessions))

    if not topic and not raw_content:
        await sio.emit(
            "airesponse",
            {"text": "❌ Please enter a topic or lesson content.", "downloadid": None, "speak": False, "language": language},
            room=sid,
        )
        return

    total_time = session_duration * num_sessions
    await sio.emit(
        "airesponse",
        {
            "text": f"📄 Generating lesson plan with worksheets...\n⏱️ {num_sessions} sessions × {session_duration} mins = {total_time} mins total",
            "downloadid": None,
            "speak": False,
            "language": language,
        },
        room=sid,
    )

    try:
        formatted = generate_lesson_from_topic(
            topic=topic,
            subject=subject,
            class_level=class_level,
            language=language,
            session_duration=session_duration,
            num_sessions=num_sessions,
        )

        lesson_data = build_lesson_data_from_formatted(
            topic=topic,
            subject=subject,
            class_level=class_level,
            language=language,
            formatted=formatted,
            session_duration=session_duration,
            num_sessions=num_sessions,
        )

        lesson_id = str(uuid.uuid4())[:8]
        lessons_store[lesson_id] = {"topic": topic, "lesson_data": lesson_data, "subject": subject}

        db = get_db()
        try:
            crud.create_lesson_plan(
                db=db,
                lesson_id=lesson_id,
                topic=topic,
                subject=subject,
                class_level=class_level,
                language=language,
                num_sessions=num_sessions,
                session_duration=session_duration,
                formatted_text=formatted,
            )
            for session in lesson_data["sessions"]:
                crud.create_lesson_session(
                    db=db,
                    lesson_id=lesson_id,
                    session_number=session["number"],
                    duration=session["duration"],
                    competency=session["competency"],
                    elo=session["elo"],
                    activities=session["activities"],
                    resources_tlm=session["resources_tlm"],
                    worksheets=session["worksheets"],
                    assessment=session["assessment"],
                )
        finally:
            db.close()

        num_worksheets = len(lesson_data.get("worksheets", []))
        num_yt_links = len(lesson_data.get("youtube_links", []))
        
        # Build worksheet summary
        worksheet_summary = ""
        for ws in lesson_data.get("worksheets", []):
            worksheet_summary += f"\n  📋 {ws.get('title', 'Worksheet')}"

        await sio.emit(
            "airesponse",
            {
                "text": (
                    "✅ LESSON PLAN WITH WORKSHEETS READY!\n\n"
                    f"📚 Topic: {topic}\n"
                    f"🎓 Class: {class_level}\n"
                    f"📖 Subject: {subject}\n"
                    f"🌐 Language: {language}\n\n"
                    f"⏱️ Duration: {num_sessions} sessions × {session_duration} mins = {total_time} mins total\n\n"
                    f"📝 Worksheets Included: {num_worksheets}{worksheet_summary}\n\n"
                    f"📺 YouTube Links: {num_yt_links}\n\n"
                    "📥 Download DOCX/PDF now to get the complete lesson plan with all worksheets!"
                ),
                "downloadid": lesson_id,
                "topic": topic,
                "speak": False,
                "language": language,
            },
            room=sid,
        )

    except Exception as e:
        logging.exception("Text generation failed")
        await sio.emit(
            "airesponse",
            {"text": f"❌ Generation failed: {str(e)}", "downloadid": None, "speak": False, "language": language},
            room=sid,
        )


@sio.event
async def voicemessage(sid, data):
    try:
        audio_b64 = data.get("audio", "")
        if not audio_b64:
            await sio.emit(
                "airesponse",
                {"text": "❌ No audio received.", "speak": True, "language": "English", "downloadid": None},
                room=sid,
            )
            return

        # Initialize conversation
        if sid not in conversation_states:
            conversation_states[sid] = ConversationState(sid)
            await sio.emit(
                "airesponse",
                {"text": PROMPTS["LANGUAGE_SELECTION"]["english"], "speak": True, "language": "English", "downloadid": None},
                room=sid,
            )
            return

        state = conversation_states[sid]

        # Decode + Transcribe
        audio_bytes = base64.b64decode(audio_b64)
        transcript = transcribe_audio_with_groq(audio_bytes)
        if not transcript:
            msg = "Could not understand. Please try again." if state.language_mode == "english" else "समझ नहीं आया। कृपया फिर से प्रयास करें।"
            await sio.emit(
                "airesponse",
                {"text": msg, "speak": True, "language": state.data.get("language", "English"), "downloadid": None},
                room=sid,
            )
            return

        state.add_to_history("user", transcript)
        parsed = parse_user_response(transcript, state.stage, state.language_mode)

        # Restart
        if parsed.get("value") == "RESTART":
            prev_lang = state.language_mode or "english"
            conversation_states[sid] = ConversationState(sid)
            conversation_states[sid].language_mode = prev_lang
            conversation_states[sid].data["language"] = prev_lang.title()

            await sio.emit(
                "airesponse",
                {"text": "Starting over..." if prev_lang == "english" else "फिर से शुरू कर रहे हैं...", "speak": True, "language": prev_lang.title(), "downloadid": None},
                room=sid,
            )
            await sio.emit(
                "airesponse",
                {"text": PROMPTS["LANGUAGE_SELECTION"]["english"], "speak": True, "language": "English", "downloadid": None},
                room=sid,
            )
            return

        response_text = None
        response_lang = state.data.get("language", "English")

        # Stage machine
        if state.stage == "LANGUAGE_SELECTION":
            if parsed["needs_clarification"]:
                response_text = PROMPTS["LANGUAGE_SELECTION"]["english"]
                response_lang = "English"
            else:
                state.language_mode = parsed["value"]
                state.data["language"] = state.language_mode.title()
                state.stage = "TOPIC_COLLECTION"
                response_text = PROMPTS["TOPIC_COLLECTION"][state.language_mode]
                response_lang = state.data["language"]

        elif state.stage == "TOPIC_COLLECTION":
            if parsed["needs_clarification"]:
                response_text = "Please tell me the lesson topic clearly." if state.language_mode == "english" else "कृपया विषय स्पष्ट रूप से बताएं।"
            else:
                state.data["topic"] = parsed["value"]
                state.stage = "SUBJECT_COLLECTION"
                response_text = PROMPTS["SUBJECT_COLLECTION"][state.language_mode]

        elif state.stage == "SUBJECT_COLLECTION":
            if parsed["needs_clarification"]:
                response_text = "Please specify the subject name." if state.language_mode == "english" else "कृपया विषय का नाम बताएं।"
            else:
                state.data["subject"] = parsed["value"]
                state.stage = "CLASS_COLLECTION"
                response_text = PROMPTS["CLASS_COLLECTION"][state.language_mode]

        elif state.stage == "CLASS_COLLECTION":
            if parsed["needs_clarification"]:
                response_text = "Please say a class number between 1 and 12." if state.language_mode == "english" else "कृपया 1 से 12 के बीच कक्षा संख्या बताएं।"
            else:
                state.data["class_level"] = parsed["value"]
                state.stage = "SESSION_DURATION"
                response_text = PROMPTS["SESSION_DURATION"][state.language_mode]

        elif state.stage == "SESSION_DURATION":
            if parsed["needs_clarification"]:
                response_text = "Please say duration between 15 and 90 minutes." if state.language_mode == "english" else "कृपया 15 से 90 मिनट के बीच अवधि बताएं।"
            else:
                state.data["session_duration"] = parsed["value"]
                state.stage = "NUM_SESSIONS"
                response_text = PROMPTS["NUM_SESSIONS"][state.language_mode]

        elif state.stage == "NUM_SESSIONS":
            if parsed["needs_clarification"]:
                response_text = "Please say number of sessions between 1 and 10." if state.language_mode == "english" else "कृपया 1 से 10 के बीच सत्रों की संख्या बताएं।"
            else:
                state.data["num_sessions"] = parsed["value"]
                state.stage = "CONFIRMATION"
                response_text = PROMPTS["CONFIRMATION"][state.language_mode](state.data)

        elif state.stage == "CONFIRMATION":
            if parsed["needs_clarification"]:
                response_text = "Please say Yes or No." if state.language_mode == "english" else "कृपया हाँ या नहीं कहें।"
            elif parsed["value"] is False:
                lang = state.language_mode or "english"
                conversation_states[sid] = ConversationState(sid)
                conversation_states[sid].language_mode = lang
                conversation_states[sid].data["language"] = lang.title()
                conversation_states[sid].stage = "TOPIC_COLLECTION"
                response_text = PROMPTS["TOPIC_COLLECTION"][lang]
            else:
                # Generate
                lang = state.data.get("language", "English")
                await sio.emit(
                    "airesponse",
                    {"text": "Generating your complete lesson plan with worksheets. Please wait.", "speak": True, "language": lang, "downloadid": None},
                    room=sid,
                )

                formatted = generate_lesson_from_topic(
                    topic=state.data["topic"],
                    subject=state.data["subject"],
                    class_level=state.data["class_level"],
                    language=state.data["language"],
                    session_duration=state.data["session_duration"],
                    num_sessions=state.data["num_sessions"],
                )
                lesson_data = build_lesson_data_from_formatted(
                    topic=state.data["topic"],
                    subject=state.data["subject"],
                    class_level=state.data["class_level"],
                    language=state.data["language"],
                    formatted=formatted,
                    session_duration=state.data["session_duration"],
                    num_sessions=state.data["num_sessions"],
                )

                lesson_id = str(uuid.uuid4())[:8]
                lessons_store[lesson_id] = {"topic": state.data["topic"], "lesson_data": lesson_data, "subject": state.data["subject"]}

                db = get_db()
                try:
                    crud.create_lesson_plan(
                        db=db,
                        lesson_id=lesson_id,
                        topic=state.data["topic"],
                        subject=state.data["subject"],
                        class_level=state.data["class_level"],
                        language=state.data["language"],
                        num_sessions=state.data["num_sessions"],
                        session_duration=state.data["session_duration"],
                        formatted_text=formatted,
                    )
                    for session in lesson_data["sessions"]:
                        crud.create_lesson_session(
                            db=db,
                            lesson_id=lesson_id,
                            session_number=session["number"],
                            duration=session["duration"],
                            competency=session["competency"],
                            elo=session["elo"],
                            activities=session["activities"],
                            resources_tlm=session["resources_tlm"],
                            worksheets=session["worksheets"],
                            assessment=session["assessment"],
                        )
                finally:
                    db.close()

                num_worksheets = len(lesson_data.get("worksheets", []))
                num_yt_links = len(lesson_data.get("youtube_links", []))

                if state.language_mode == "hindi":
                    success_msg = (
                        f"पाठ योजना तैयार है! विषय: {state.data['topic']}। "
                        f"कक्षा: {state.data['class_level']}। "
                        f"विषय: {state.data['subject']}। "
                        f"{num_worksheets} वर्कशीट और {num_yt_links} YouTube लिंक तैयार हैं। "
                        "अब आप डाउनलोड कर सकते हैं।"
                    )
                else:
                    success_msg = (
                        f"Lesson plan ready! Topic: {state.data['topic']}. "
                        f"Class: {state.data['class_level']}. "
                        f"Subject: {state.data['subject']}. "
                        f"{num_worksheets} worksheets and {num_yt_links} YouTube links are ready. "
                        "You can now download the complete lesson plan with all worksheets."
                    )

                await sio.emit(
                    "airesponse",
                    {"text": success_msg, "speak": True, "language": state.data["language"], "downloadid": lesson_id, "topic": state.data["topic"]},
                    room=sid,
                )

                del conversation_states[sid]
                return

        if response_text:
            await sio.emit(
                "airesponse",
                {"text": response_text, "speak": True, "language": response_lang, "downloadid": None},
                room=sid,
            )

    except Exception as e:
        logging.exception("Voice conversation failed")
        await sio.emit(
            "airesponse",
            {"text": f"Error: {str(e)}", "speak": True, "language": "English", "downloadid": None},
            room=sid,
        )


# ============================================================
# Main
# ============================================================

if __name__ == "__main__":
    import uvicorn

    print("=" * 60)
    print("🚀 NCERT Lesson Plan Generator with Worksheets")
    print("=" * 60)
    print("🌐 Open: http://localhost:5000")
    print("🔧 MCP Health: http://localhost:5000/mcp/health")
    print("=" * 60)

    uvicorn.run("main:socket_app", host="0.0.0.0", port=5000, reload=True)